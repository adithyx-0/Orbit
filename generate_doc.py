from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Title
title = doc.add_heading('Orbit — Feature Sprint Work Division', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Project: AI-Powered Personal Resource Optimization & Subscription Intelligence System')
doc.add_paragraph('Course: 22AIE311 — Software Engineering\n')

# ─── Adithyan Section ───
doc.add_heading('Adithyan R (Roll: AM.SC.U4AIE23065) — Backend & AI', level=1)

adithyan_tasks = [
    ('1', 'Gamification Database', 'New tables: user_scores, achievements, streaks — to track user progress and rewards'),
    ('2', 'Star / Achievement API', 'Endpoints to award stars when goals are completed, track daily streaks, and calculate user level'),
    ('3', 'AI Chatbot Endpoint', 'POST /api/chat — integrates Gemini / Claude API with the user\'s real subscription and usage data as context'),
    ('4', 'Smart Recommendations', 'Algorithm to flag subscriptions with low usage hours and suggest cancellation with cost savings shown'),
    ('5', 'Renewal Alerts', 'SendGrid email notifications when a subscription renews within 3 days'),
    ('6', 'Usage Analytics', 'Weekly/monthly cost trends, forecast next month\'s spend, cost-per-hour ROI per subscription'),
]

table1 = doc.add_table(rows=1, cols=3)
table1.style = 'Table Grid'
hdr = table1.rows[0].cells
hdr[0].text = '#'
hdr[1].text = 'Task'
hdr[2].text = 'Details'
for cell in hdr:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True

for num, task, detail in adithyan_tasks:
    row = table1.add_row().cells
    row[0].text = num
    row[1].text = task
    row[2].text = detail

doc.add_paragraph('')

# ─── Adithya Section ───
doc.add_heading('Adithya Rajesh (Roll: AM.SC.U4AIE23007) — Frontend & UI', level=1)

adithya_tasks = [
    ('1', 'Full UI Redesign', 'Dark/gradient theme, animated cards, improved typography and layout across all pages'),
    ('2', 'Star Tree Component', 'Visual tree that grows a new star each time a goal is completed — the centerpiece gamification feature'),
    ('3', 'User Score Widget', 'Score badge on dashboard with level system: Bronze → Silver → Gold → Platinum'),
    ('4', 'AI Chat Widget', 'Floating chat bubble, sliding chat panel connected to the backend chatbot API'),
    ('5', 'Subscription Cards', 'Richer cards with app logos, usage progress bars, ROI badge, color-coded by category'),
    ('6', 'Onboarding Wizard', 'Step-by-step first-time setup — add subscriptions, set goals, pick interests'),
    ('7', 'Mobile Responsive', 'Fully responsive layout for all phone and tablet screen sizes'),
]

table2 = doc.add_table(rows=1, cols=3)
table2.style = 'Table Grid'
hdr2 = table2.rows[0].cells
hdr2[0].text = '#'
hdr2[1].text = 'Task'
hdr2[2].text = 'Details'
for cell in hdr2:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True

for num, task, detail in adithya_tasks:
    row = table2.add_row().cells
    row[0].text = num
    row[1].text = task
    row[2].text = detail

doc.add_paragraph('')

# ─── Priority Order ───
doc.add_heading('Suggested Priority Order', level=1)
weeks = [
    ('Week 1', 'UI Redesign + Star Tree UI + Gamification backend tables & APIs'),
    ('Week 2', 'AI Chatbot — backend endpoint + frontend chat widget'),
    ('Week 3', 'Smart Recommendations + Renewal Email Alerts'),
    ('Week 4', 'Polish, mobile responsiveness, onboarding wizard'),
]
for week, desc in weeks:
    p = doc.add_paragraph()
    p.add_run(f'{week}: ').bold = True
    p.add_run(desc)

doc.add_paragraph('')

# ─── Novel Features ───
doc.add_heading('What Makes This App Stand Out', level=1)
novel = [
    'Star Tree — visual, emotional reward system unlike any existing subscription tracker',
    'AI Chatbot with real context — not a generic bot, it knows YOUR subscriptions and usage data',
    'ROI Score per subscription — ranks subs by value delivered per hour spent',
    'User Level System — gamified engagement that keeps users coming back',
    'Cost Forecasting — predicts next month\'s spend based on current subscriptions',
]
for point in novel:
    doc.add_paragraph(point, style='List Bullet')

doc.save('D:\\Orbit\\Work_Division.docx')
print('Done: D:\\Orbit\\Work_Division.docx')
